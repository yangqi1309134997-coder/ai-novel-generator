"""
文件解析模块 - 支持 txt/pdf/epub，带进度跟踪和错误处理，支持自定义章节模板

版权所有 © 2026 新疆幻城网安科技有限责任公司 (幻城科技)
作者：幻城
"""
import os
import re
import logging
import tempfile
from typing import Tuple, List, Optional, IO, Dict
from enum import Enum
from dataclasses import dataclass

logger = logging.getLogger(__name__)

# 常量
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
MIN_PARAGRAPH_LENGTH = 20  # 最小段落长度


# 预设章节模板
CHAPTER_PATTERNS = {
    "默认": [
        r'第\s*\d+\s*章[：:\s]*.*',
        r'第\s*\d+\s*章',
        r'Chapter\s*\d+',
    ],
    "简洁格式": [
        r'^\d+\.',
        r'^\d+、',
        r'^\d+\s',
    ],
    "带书名号": [
        r'《第\d+章》',
        r'「第\d+章」',
    ],
    "英文格式": [
        r'Chapter\s+\d+[:：\s]*.*',
        r'CHAPTER\s+\d+[:：\s]*.*',
        r'Part\s+\d+',
    ],
    "特殊格式": [
        r'【.*第\d+章.*】',
        r'≮.*第\d+章.*≯',
        r'◆.*第\d+章.*◆',
    ],
}


@dataclass
class ChapterInfo:
    """章节信息"""
    num: int
    title: str
    content: str
    start_pos: int = 0
    end_pos: int = 0


class FileType(Enum):
    """文件类型"""
    TXT = "txt"
    PDF = "pdf"
    EPUB = "epub"
    MD = "md"
    DOCX = "docx"
    UNKNOWN = "unknown"


def get_file_type(file_path: str) -> FileType:
    """获取文件类型"""
    if not file_path:
        return FileType.UNKNOWN
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".txt":
        return FileType.TXT
    elif ext == ".pdf":
        return FileType.PDF
    elif ext == ".epub":
        return FileType.EPUB
    elif ext == ".md":
        return FileType.MD
    elif ext == ".docx":
        return FileType.DOCX
    else:
        return FileType.UNKNOWN


def parse_txt_file(file_path: str) -> Tuple[List[str], str]:
    """
    解析TXT文件
    
    Returns:
        (段落列表, 状态信息)
    """
    try:
        # 支持传入文件对象或路径
        if hasattr(file_path, 'read'):
            fobj: IO = file_path
            # 尝试获取 size 属性
            try:
                fobj.seek(0, os.SEEK_END)
                file_size = fobj.tell()
                fobj.seek(0)
            except Exception:
                file_size = 0
        else:
            file_size = os.path.getsize(file_path)

        if file_size and file_size > MAX_FILE_SIZE:
            return [], f"错误：文件过大 ({file_size / 1024 / 1024:.1f}MB > 50MB)"

        paragraphs: List[str] = []
        buf_lines: List[str] = []
        total_chars = 0

        # 逐行读取以降低内存压力
        if hasattr(file_path, 'read'):
            stream = file_path
        else:
            stream = open(file_path, 'r', encoding='utf-8', errors='ignore')

        try:
            for line in stream:
                stripped = line.rstrip('\n')
                total_chars += len(stripped)

                if stripped.strip() == '':
                    # 空行 -> 段落结束
                    if buf_lines:
                        para = '\n'.join(buf_lines).strip()
                        if len(para) >= MIN_PARAGRAPH_LENGTH:
                            paragraphs.append(para)
                        buf_lines = []
                    continue

                # 常规行
                buf_lines.append(stripped)

            # 最后一段
            if buf_lines:
                para = '\n'.join(buf_lines).strip()
                if len(para) >= MIN_PARAGRAPH_LENGTH:
                    paragraphs.append(para)

        finally:
            if not hasattr(file_path, 'read'):
                stream.close()

        logger.info(f"TXT文件解析完成，共 {len(paragraphs)} 段")
        return paragraphs, f"解析完成，共 {len(paragraphs)} 段，约 {total_chars} 字"
    
    except Exception as e:
        logger.error(f"TXT文件解析失败: {e}")
        return [], f"读取失败：{str(e)}"


def parse_pdf_file(file_path: str) -> Tuple[List[str], str]:
    """
    解析PDF文件
    
    Returns:
        (段落列表, 状态信息)
    """
    try:
        import fitz
    except ImportError:
        return [], "错误：缺少PyMuPDF依赖，请运行: pip install PyMuPDF"
    
    try:
        file_size = os.path.getsize(file_path)
        if file_size > MAX_FILE_SIZE:
            return [], f"错误：文件过大 ({file_size / 1024 / 1024:.1f}MB > 50MB)"
        
        text = ""
        doc = fitz.open(file_path)
        
        for page_num, page in enumerate(doc):
            try:
                page_text = page.get_text("text")
                text += page_text + "\n"
            except Exception as e:
                logger.warning(f"PDF页面 {page_num} 解析失败: {e}")
        
        doc.close()
        
        paragraphs = _split_paragraphs(text)
        logger.info(f"PDF文件解析完成，共 {len(paragraphs)} 段")
        return paragraphs, f"解析完成，共 {len(paragraphs)} 段，约 {len(text)} 字"
    
    except Exception as e:
        logger.error(f"PDF文件解析失败: {e}")
        return [], f"读取失败：{str(e)}"


def parse_epub_file(file_path: str) -> Tuple[List[str], str]:
    """
    解析EPUB文件
    
    Returns:
        (段落列表, 状态信息)
    """
    try:
        from ebooklib import epub
        from bs4 import BeautifulSoup
    except ImportError:
        return [], "错误：缺少ebooklib或beautifulsoup4依赖，请运行: pip install ebooklib beautifulsoup4"
    
    try:
        file_size = os.path.getsize(file_path)
        if file_size > MAX_FILE_SIZE:
            return [], f"错误：文件过大 ({file_size / 1024 / 1024:.1f}MB > 50MB)"
        
        text = ""
        book = epub.read_epub(file_path)
        
        for item in book.get_items():
            if item.get_type() == epub.ITEM_DOCUMENT:
                try:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text += soup.get_text(separator="\n") + "\n"
                except Exception as e:
                    logger.warning(f"EPUB章节解析失败: {e}")
        
        paragraphs = _split_paragraphs(text)
        logger.info(f"EPUB文件解析完成，共 {len(paragraphs)} 段")
        return paragraphs, f"解析完成，共 {len(paragraphs)} 段，约 {len(text)} 字"
    
    except Exception as e:
        logger.error(f"EPUB文件解析失败: {e}")
        return [], f"读取失败：{str(e)}"


def parse_md_file(file_path: str) -> Tuple[List[str], str]:
    """
    解析Markdown文件
    
    Returns:
        (段落列表, 状态信息)
    """
    try:
        # 支持传入文件对象或路径
        if hasattr(file_path, 'read'):
            fobj: IO = file_path
            # 尝试获取 size 属性
            try:
                fobj.seek(0, os.SEEK_END)
                file_size = fobj.tell()
                fobj.seek(0)
            except Exception:
                file_size = 0
        else:
            file_size = os.path.getsize(file_path)

        if file_size and file_size > MAX_FILE_SIZE:
            return [], f"错误：文件过大 ({file_size / 1024 / 1024:.1f}MB > 50MB)"

        paragraphs: List[str] = []
        buf_lines: List[str] = []
        total_chars = 0

        # 逐行读取以降低内存压力
        if hasattr(file_path, 'read'):
            stream = file_path
        else:
            stream = open(file_path, 'r', encoding='utf-8', errors='ignore')

        try:
            for line in stream:
                stripped = line.rstrip('\n')
                total_chars += len(stripped)

                if stripped.strip() == '':
                    # 空行 -> 段落结束
                    if buf_lines:
                        para = '\n'.join(buf_lines).strip()
                        if len(para) >= MIN_PARAGRAPH_LENGTH:
                            paragraphs.append(para)
                        buf_lines = []
                    continue

                # 常规行
                buf_lines.append(stripped)

            # 最后一段
            if buf_lines:
                para = '\n'.join(buf_lines).strip()
                if len(para) >= MIN_PARAGRAPH_LENGTH:
                    paragraphs.append(para)

        finally:
            if not hasattr(file_path, 'read'):
                stream.close()

        logger.info(f"Markdown文件解析完成，共 {len(paragraphs)} 段")
        return paragraphs, f"解析完成，共 {len(paragraphs)} 段，约 {total_chars} 字"
    
    except Exception as e:
        logger.error(f"Markdown文件解析失败: {e}")
        return [], f"读取失败：{str(e)}"


def parse_docx_file(file_path: str) -> Tuple[List[str], str]:
    """
    解析Word文档文件
    
    Returns:
        (段落列表, 状态信息)
    """
    try:
        from docx import Document
    except ImportError:
        return [], "错误：缺少python-docx依赖，请运行: pip install python-docx"
    
    try:
        file_size = os.path.getsize(file_path)
        if file_size > MAX_FILE_SIZE:
            return [], f"错误：文件过大 ({file_size / 1024 / 1024:.1f}MB > 50MB)"
        
        doc = Document(file_path)
        paragraphs: List[str] = []
        total_chars = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) >= MIN_PARAGRAPH_LENGTH:
                paragraphs.append(text)
                total_chars += len(text)
        
        logger.info(f"Word文档解析完成，共 {len(paragraphs)} 段")
        return paragraphs, f"解析完成，共 {len(paragraphs)} 段，约 {total_chars} 字"
    
    except Exception as e:
        logger.error(f"Word文档解析失败: {e}")
        return [], f"读取失败：{str(e)}"


def parse_novel_file(file_path: str) -> Tuple[List[str], str]:
    """
    解析小说文件（自动识别格式）
    
    Args:
        file_path: 文件路径
    
    Returns:
        (段落列表, 状态信息)
    """
    if not file_path:
        return [], "无文件"
    
    # 处理Gradio上传的文件对象或文件流
    temp_path = None
    if hasattr(file_path, 'name') and isinstance(file_path.name, str) and os.path.exists(file_path.name):
        file_path = file_path.name
    elif hasattr(file_path, 'read'):
        # 将上传的流写入临时文件以便下游库处理（PDF/EPUB/DOCX 需要文件路径）
        try:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.tmp')
            chunk = file_path.read(8192)
            while chunk:
                if isinstance(chunk, str):
                    tmp.write(chunk.encode('utf-8'))
                else:
                    tmp.write(chunk)
                chunk = file_path.read(8192)
            tmp.close()
            temp_path = tmp.name
            file_path = temp_path
        except Exception as e:
            logger.error(f"处理上传文件失败: {e}")
            return [], f"读取上传文件失败: {e}"
    
    if not os.path.exists(file_path):
        return [], f"文件不存在: {file_path}"
    
    file_type = get_file_type(file_path)
    
    if file_type == FileType.TXT:
        try:
            return parse_txt_file(file_path)
        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass
    elif file_type == FileType.PDF:
        try:
            return parse_pdf_file(file_path)
        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass
    elif file_type == FileType.EPUB:
        try:
            return parse_epub_file(file_path)
        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass
    elif file_type == FileType.MD:
        try:
            return parse_md_file(file_path)
        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass
    elif file_type == FileType.DOCX:
        try:
            return parse_docx_file(file_path)
        finally:
            if temp_path:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass
    else:
        return [], "不支持的文件格式（支持 txt/pdf/epub/md/docx）"


def _split_paragraphs(text: str, min_length: int = MIN_PARAGRAPH_LENGTH) -> List[str]:
    """
    将文本分割为段落
    
    Args:
        text: 原始文本
        min_length: 最小段落长度
    
    Returns:
        段落列表
    """
    # 按多个换行符分割
    raw_paragraphs = re.split(r'\n\s*\n+', text)
    
    # 清理和过滤
    paragraphs = []
    for para in raw_paragraphs:
        para = para.strip()
        # 移除章节标题等特殊标记
        para = re.sub(r'^(第\d+章|Chapter \d+|第 \d+ 章)[：:：]?\s*', '', para)
        para = re.sub(r'^\s*\*+\s*|\s*\*+\s*$', '', para)
        
        if len(para) >= min_length:
            paragraphs.append(para)
    
    return paragraphs


def estimate_word_count(text: str) -> int:
    """估计中文字数（粗略估计）"""
    chinese_count = len(re.findall(r'[\u4e00-\u9fff]', text))
    english_count = len(re.findall(r'\b[a-zA-Z]+\b', text))
    # 中文按1字计算，英文按0.5字计算
    return chinese_count + int(english_count * 0.5)


def parse_novel_by_chapters(
    file_path: str,
    pattern_name: str = "默认",
    custom_pattern: str = ""
) -> Tuple[List[ChapterInfo], str]:
    """
    按章节解析小说文件

    Args:
        file_path: 文件路径
        pattern_name: 预设模板名称
        custom_pattern: 自定义正则表达式（如果提供，则优先使用）

    Returns:
        (章节列表, 状态信息)
    """
    try:
        # 读取文本
        file_type = get_file_type(file_path)

        if file_type == FileType.TXT:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        elif file_type == FileType.PDF:
            import fitz
            text = ""
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text("text") + "\n"
            doc.close()
        elif file_type == FileType.EPUB:
            from ebooklib import epub
            from bs4 import BeautifulSoup
            text = ""
            book = epub.read_epub(file_path)
            for item in book.get_items():
                if item.get_type() == epub.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text += soup.get_text(separator="\n") + "\n"
        elif file_type == FileType.MD:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        elif file_type == FileType.DOCX:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            return [], "不支持的文件格式"

        # 确定使用的正则表达式
        if custom_pattern and custom_pattern.strip():
            patterns = [custom_pattern.strip()]
        elif pattern_name in CHAPTER_PATTERNS:
            patterns = CHAPTER_PATTERNS[pattern_name]
        else:
            patterns = CHAPTER_PATTERNS["默认"]

        # 查找所有章节标题
        chapters = []
        lines = text.split('\n')

        current_chapter_num = 0
        current_chapter_title = ""
        current_chapter_content = []
        chapter_start_pos = 0

        for i, line in enumerate(lines):
            line_stripped = line.strip()
            is_chapter_header = False

            # 检查是否匹配任何章节模式
            for pattern in patterns:
                if re.match(pattern, line_stripped, re.IGNORECASE):
                    is_chapter_header = True
                    break

            if is_chapter_header:
                # 保存上一章
                if current_chapter_num > 0:
                    content = '\n'.join(current_chapter_content).strip()
                    if content:
                        chapters.append(ChapterInfo(
                            num=current_chapter_num,
                            title=current_chapter_title,
                            content=content,
                            start_pos=chapter_start_pos,
                            end_pos=i
                        ))

                # 提取章节号和标题
                current_chapter_num += 1
                current_chapter_title = line_stripped
                current_chapter_content = []
                chapter_start_pos = i
            else:
                # 跳过空行但保留内容
                if line_stripped or current_chapter_content:
                    current_chapter_content.append(line)

        # 保存最后一章
        if current_chapter_num > 0 and current_chapter_content:
            content = '\n'.join(current_chapter_content).strip()
            if content:
                chapters.append(ChapterInfo(
                    num=current_chapter_num,
                    title=current_chapter_title,
                    content=content,
                    start_pos=chapter_start_pos,
                    end_pos=len(lines)
                ))

        logger.info(f"按章节解析完成，共 {len(chapters)} 章")
        return chapters, f"解析完成，共 {len(chapters)} 章"

    except Exception as e:
        logger.error(f"按章节解析失败: {e}")
        return [], f"解析失败: {str(e)}"


def parse_novel_with_custom_template(
    file_path: str,
    custom_template: str
) -> Tuple[List[ChapterInfo], str]:
    """
    使用自定义模板解析小说

    Args:
        file_path: 文件路径
        custom_template: 自定义章节模板（支持占位符）
                       例如: "第{n}章 {title}" 或 "Chapter {n}: {title}"

    Returns:
        (章节列表, 状态信息)
    """
    if not custom_template or not custom_template.strip():
        return parse_novel_by_chapters(file_path, "默认", "")

    # 将模板转换为正则表达式
    # {n} 或 {num} -> (\d+)
    # {title} -> (.*)
    pattern = custom_template.strip()
    pattern = re.escape(pattern)
    pattern = pattern.replace(r'\{n\}', r'(\d+)')
    pattern = pattern.replace(r'\{num\}', r'(\d+)')
    pattern = pattern.replace(r'\{title\}', r'(.*)')
    pattern = pattern.replace(r'\{.*?\}', r'.*')  # 其他占位符

    # 确保匹配行首
    if not pattern.startswith('^'):
        pattern = '^' + pattern

    return parse_novel_by_chapters(file_path, custom_pattern=pattern)


def split_by_word_count(text: str, word_count: int) -> List[str]:
    """
    按字数分段

    Args:
        text: 原始文本
        word_count: 每段的字数

    Returns:
        分段后的文本列表
    """
    if not text or not text.strip():
        return []

    if word_count <= 0:
        raise ValueError("字数必须大于0")

    # 按字数均匀分段
    segments = []
    total_length = len(text)
    start = 0

    while start < total_length:
        end = start + word_count
        if end > total_length:
            end = total_length

        segment = text[start:end].strip()
        if segment:
            segments.append(segment)

        start = end

    logger.info(f"按字数分段完成，共 {len(segments)} 段，每段约 {word_count} 字")
    return segments


def split_by_pattern(text: str, pattern: str, keep_marker: bool = True) -> List[str]:
    """
    按固定文本/变量分段

    Args:
        text: 原始文本
        pattern: 分段标记（支持变量：%章、%节、%回，或自定义文本）
        keep_marker: 是否保留分段标记

    Returns:
        分段后的文本列表
    """
    if not text or not text.strip():
        return []

    if not pattern or not pattern.strip():
        raise ValueError("分段标记不能为空")

    # 智能识别：如果用户输入"第x章"、"第X章"等，自动转换为正则表达式
    # 检查是否包含"第"和"章"、"节"、"回"的组合
    pattern_lower = pattern.strip().lower()
    
    # 检查是否是简化的模式（如"第x章"、"第X章"）
    if pattern_lower in ['第x章', '第x章', '第x章', '第x章']:
        # 同时支持中文数字和阿拉伯数字，使用+确保至少匹配一个数字
        # 使用负向预查确保"第x章"后面不能有中文字符（除了空格和标点符号）
        # 匹配格式：第x章、第x章:、第x章：、第x章 (空格)、第x章 (空格后换行)
        # 支持Markdown格式：## 第x章
        # 但不匹配：第一章这是、第一章的内容等（后面有中文字符）
        regex_pattern = r'^[\s#]*\s*第\s*[一二三四五六七八九十百千万零〇0123456789]+\s*章\s*[:：\s]*(?![\u4e00-\u9fff])'
        logger.info("检测到'第x章'模式，自动转换为正则表达式")
    elif pattern_lower in ['第x节', '第x节', '第x节', '第x节']:
        regex_pattern = r'^\s*第\s*[一二三四五六七八九十百千万零〇0123456789]+\s*节\s*[:：\s]*(?![\u4e00-\u9fff])'
        logger.info("检测到'第x节'模式，自动转换为正则表达式")
    elif pattern_lower in ['第x回', '第x回', '第x回', '第x回']:
        regex_pattern = r'^\s*第\s*[一二三四五六七八九十百千万零〇0123456789]+\s*回\s*[:：\s]*(?![\u4e00-\u9fff])'
        logger.info("检测到'第x回'模式，自动转换为正则表达式")
    elif '%章' in pattern_lower or '%节' in pattern_lower or '%回' in pattern_lower:
        # 使用变量替换
        regex_pattern = pattern.strip()
        # %章 -> 匹配"第X章"、"第x章"等（支持中文和阿拉伯数字）
        regex_pattern = regex_pattern.replace('%章', r'[一二三四五六七八九十百千万零〇0123456789]+\s*章')
        # %节 -> 匹配"第X节"、"第x节"等（支持中文和阿拉伯数字）
        regex_pattern = regex_pattern.replace('%节', r'[一二三四五六七八九十百千万零〇0123456789]+\s*节')
        # %回 -> 匹配"第X回"、"第x回"等（支持中文和阿拉伯数字）
        regex_pattern = regex_pattern.replace('%回', r'[一二三四五六七八九十百千万零〇0123456789]+\s*回')
        # 确保正则表达式以^开头（匹配行首）
        if not regex_pattern.startswith('^'):
            regex_pattern = '^' + regex_pattern
    else:
        # 不包含章节标记，直接使用原始模式
        regex_pattern = pattern.strip()

    # 尝试按模式分割
    try:
        # 如果保留标记，使用正则表达式查找所有匹配位置
        if keep_marker:
            # 查找所有匹配的位置
            matches = list(re.finditer(regex_pattern, text, flags=re.MULTILINE | re.IGNORECASE))

            if not matches:
                # 没有匹配，返回整个文本
                logger.warning(f"未找到匹配的模式: {regex_pattern}，返回整个文本")
                return [text.strip()] if text.strip() else []

            segments = []
            prev_end = 0

            for match in matches:
                # 获取匹配的标记
                marker = match.group(0)

                # 获取标记之前的内容（如果有）
                if prev_end < match.start():
                    prev_content = text[prev_end:match.start()].strip()
                    if prev_content:
                        segments.append(prev_content)

                # 添加标记
                segments.append(marker.strip())

                prev_end = match.end()

            # 添加最后一段内容
            if prev_end < len(text):
                last_content = text[prev_end:].strip()
                if last_content:
                    segments.append(last_content)

            # 合并标记和内容
            result = []
            i = 0
            while i < len(segments):
                # 如果当前是标记，且后面有内容
                if i + 1 < len(segments):
                    result.append((segments[i] + segments[i + 1]).strip())
                    i += 2
                else:
                    # 只有标记或内容
                    if segments[i].strip():
                        result.append(segments[i].strip())
                    i += 1

            segments = result
        else:
            # 不保留标记，直接分割
            segments = re.split(regex_pattern, text, flags=re.MULTILINE | re.IGNORECASE)

        # 清理空段落
        segments = [seg.strip() for seg in segments if seg.strip()]

        logger.info(f"按模式分段完成，共 {len(segments)} 段")
        return segments

    except re.error as e:
        raise ValueError(f"无效的正则表达式: {e}")
