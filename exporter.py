"""
导出模块 - 支持 Word (DOCX)、TXT、Markdown

版权所有 © 2026 新疆幻城网安科技有限责任公司 (幻城科技)
作者：幻城
"""
import os
import re
import logging
import tempfile
from typing import Tuple, Optional
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)

MODULE_ROOT = os.path.dirname(os.path.abspath(__file__))
EXPORT_DIR = os.path.join(MODULE_ROOT, "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)


def _sanitize_filename(name: str, max_len: int = 120) -> str:
    """清理文件名中的非法字符并限制长度"""
    if not name or not name.strip():
        name = "novel"
    safe = re.sub(r'[<>:"/\\|?*]', '_', name).strip()
    if len(safe) > max_len:
        safe = safe[:max_len]
    return safe


def _extract_chapters_from_markdown(text: str) -> list:
    """
    从Markdown格式的小说文本中提取章节信息
    
    Returns:
        [{"title": "...", "content": "..."}, ...]
    """
    chapters = []
    current_chapter = None
    content_lines = []

    # 更通用的章节头检测，支持 '#', '##', '###' 等级别，支持中文与空格变体
    header_re = re.compile(r'^(?:#{1,6}\s*)?(第\s*\d+\s*章[\s\S]*|第\s*\d+章[\s\S]*|第\s*\d+\s*章[:：\s\-—–]?.*)$', re.IGNORECASE)

    for line in text.splitlines():
        if not line:
            # 空行作为段落分隔符，但不要结束章节
            if current_chapter:
                content_lines.append('')
            continue

        # 检测章节标题
        if header_re.match(line.strip()):
            # 保存上一章
            if current_chapter:
                current_chapter['content'] = '\n'.join([l for l in content_lines]).strip()
                chapters.append(current_chapter)

            # 提取标题文本
            title_match = re.search(r'(第\s*\d+\s*章[\s\S]*)', line)
            title = title_match.group(1).strip() if title_match else line.strip()
            current_chapter = {'title': title, 'content': ''}
            content_lines = []
            continue

        # 跳过文件级标题
        if line.strip().startswith('# '):
            continue

        if current_chapter is None:
            # 如果还未遇到章节头，把内容作为第一章的一部分
            current_chapter = {'title': '第一章', 'content': ''}
            content_lines = [line]
        else:
            content_lines.append(line)

    # 保存最后一章
    if current_chapter:
        current_chapter['content'] = '\n'.join([l for l in content_lines]).strip()
        chapters.append(current_chapter)

    return chapters


def export_to_txt(novel_text: str, title: str) -> Tuple[Optional[str], str]:
    """
    导出为TXT格式
    
    Args:
        novel_text: 小说文本（Markdown格式）
        title: 小说标题
    
    Returns:
        (文件路径, 状态信息)
    """
    try:
        if not novel_text.strip():
            return None, "无内容可导出"
        
        # 提取章节
        chapters = _extract_chapters_from_markdown(novel_text)
        
        if not chapters:
            return None, "无法从文本中提取章节"
        
        # 生成TXT内容
        txt_content = f"{title}\n\n"
        
        for chapter in chapters:
            txt_content += f"{chapter['title']}\n\n"
            txt_content += f"{chapter['content']}\n\n"
            txt_content += "-" * 80 + "\n\n"
        
        # 保存文件（原子写入）
        safe_title = _sanitize_filename(title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        filepath = os.path.join(EXPORT_DIR, filename)
        try:
            with tempfile.NamedTemporaryFile('w', encoding='utf-8', delete=False, dir=EXPORT_DIR) as tmp:
                tmp.write(txt_content)
                tmp_path = tmp.name
            os.replace(tmp_path, filepath)
        except Exception as e:
            logger.error(f"写入TXT文件失败: {e}")
            return None, f"导出失败: {e}"

        logger.info(f"TXT导出成功: {filename}")
        return filepath, f"导出成功: {filename}"
    
    except Exception as e:
        logger.error(f"TXT导出失败: {e}")
        return None, f"导出失败: {str(e)}"


def export_to_markdown(novel_text: str, title: str) -> Tuple[Optional[str], str]:
    """
    导出为Markdown格式
    
    Args:
        novel_text: 小说文本（Markdown格式）
        title: 小说标题
    
    Returns:
        (文件路径, 状态信息)
    """
    try:
        if not novel_text.strip():
            return None, "无内容可导出"
        
        # 添加元数据
        md_content = f"# {title}\n\n"
        md_content += f"*生成于: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n"
        md_content += "---\n\n"
        md_content += novel_text
        
        # 保存文件（原子写入）
        safe_title = _sanitize_filename(title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        filepath = os.path.join(EXPORT_DIR, filename)
        try:
            with tempfile.NamedTemporaryFile('w', encoding='utf-8', delete=False, dir=EXPORT_DIR) as tmp:
                tmp.write(md_content)
                tmp_path = tmp.name
            os.replace(tmp_path, filepath)
        except Exception as e:
            logger.error(f"写入Markdown文件失败: {e}")
            return None, f"导出失败: {e}"

        logger.info(f"Markdown导出成功: {filename}")
        return filepath, f"导出成功: {filename}"
    
    except Exception as e:
        logger.error(f"Markdown导出失败: {e}")
        return None, f"导出失败: {str(e)}"


def export_to_docx(novel_text: str, title: str) -> Tuple[Optional[str], str]:
    """
    导出为Word (DOCX) 格式 - 专业排版
    
    Args:
        novel_text: 小说文本（Markdown格式）
        title: 小说标题
    
    Returns:
        (文件路径, 状态信息)
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        return None, "错误：缺少python-docx依赖，请运行: pip install python-docx"
    
    try:
        if not novel_text.strip():
            return None, "无内容可导出"
        
        # 提取章节
        chapters = _extract_chapters_from_markdown(novel_text)
        
        if not chapters:
            return None, "无法从文本中提取章节"
        
        doc = Document()
        
        # 配置样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        # 中文字体
        rPr = style.element.get_or_add_rPr()
        rPr.find(qn('w:rFonts')).set(qn('w:eastAsia'), '宋体')
        
        # 段落格式
        style.paragraph_format.first_line_indent = Pt(24)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.5
        
        # 添加书名
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.name = '黑体'
        title_run.font.size = Pt(26)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 中文字体设置
        title_rPr = title_run._element.get_or_add_rPr()
        title_rPr.find(qn('w:rFonts')).set(qn('w:eastAsia'), '黑体')
        
        # 添加作者和日期信息
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
        info_run.font.size = Pt(10)
        
        doc.add_paragraph()  # 空行
        
        # 添加章节
        for chapter in chapters:
            # 章节标题
            chapter_title_para = doc.add_paragraph(chapter['title'])
            chapter_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for run in chapter_title_para.runs:
                run.font.name = '黑体'
                run.font.size = Pt(16)
                run.font.bold = True
                run_rPr = run._element.get_or_add_rPr()
                run_rPr.find(qn('w:rFonts')).set(qn('w:eastAsia'), '黑体')
            
            doc.add_paragraph()  # 空行
            
            # 章节内容 - 按段落添加
            paragraphs = chapter['content'].split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip(), style='Normal')
            
            doc.add_paragraph()  # 章节间空行
        
        # 保存文件（原子写入）
        safe_title = _sanitize_filename(title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(EXPORT_DIR, filename)
        try:
            tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx', dir=EXPORT_DIR)
            os.close(tmp_fd)
            doc.save(tmp_path)
            os.replace(tmp_path, filepath)
        except Exception as e:
            logger.error(f"写入DOCX文件失败: {e}")
            return None, f"导出失败: {e}"

        logger.info(f"DOCX导出成功: {filename}")
        return filepath, f"导出成功: {filename}"
    
    except Exception as e:
        logger.error(f"DOCX导出失败: {e}")
        return None, f"导出失败: {str(e)}"


def export_to_html(novel_text: str, title: str) -> Tuple[Optional[str], str]:
    """
    导出为HTML格式 - 可在浏览器中阅读
    
    Args:
        novel_text: 小说文本（Markdown格式）
        title: 小说标题
    
    Returns:
        (文件路径, 状态信息)
    """
    try:
        import markdown
    except ImportError:
        return None, "错误：缺少markdown依赖，请运行: pip install markdown"
    
    try:
        if not novel_text.strip():
            return None, "无内容可导出"
        
        # 转换Markdown为HTML
        html_content = markdown.markdown(novel_text)
        
        # 包裹为完整HTML文档
        full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Microsoft YaHei', '宋体', serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.8;
            background-color: #f5f5f5;
            color: #333;
        }}
        h1 {{
            text-align: center;
            font-size: 2.5em;
            margin-bottom: 0.5em;
        }}
        h2 {{
            text-align: center;
            font-size: 1.5em;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
            border-bottom: 2px solid #ddd;
            padding-bottom: 0.3em;
        }}
        p {{
            text-align: justify;
            text-indent: 2em;
            margin: 1em 0;
        }}
        .info {{
            text-align: center;
            color: #999;
            font-size: 0.9em;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p class="info">生成于: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    <hr>
    {html_content}
</body>
</html>"""
        
        # 保存文件（原子写入）
        safe_title = _sanitize_filename(title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        filepath = os.path.join(EXPORT_DIR, filename)
        try:
            with tempfile.NamedTemporaryFile('w', encoding='utf-8', delete=False, dir=EXPORT_DIR) as tmp:
                tmp.write(full_html)
                tmp_path = tmp.name
            os.replace(tmp_path, filepath)
        except Exception as e:
            logger.error(f"写入HTML文件失败: {e}")
            return None, f"导出失败: {e}"

        logger.info(f"HTML导出成功: {filename}")
        return filepath, f"导出成功: {filename}"
    
    except Exception as e:
        logger.error(f"HTML导出失败: {e}")
        return None, f"导出失败: {str(e)}"


def list_export_files() -> list:
    """列出所有导出文件"""
    try:
        files = []
        for filename in os.listdir(EXPORT_DIR):
            filepath = os.path.join(EXPORT_DIR, filename)
            if os.path.isfile(filepath):
                file_size = os.path.getsize(filepath)
                file_time = datetime.fromtimestamp(os.path.getmtime(filepath))
                files.append({
                    'name': filename,
                    'path': filepath,
                    'size': file_size,
                    'time': file_time.strftime('%Y-%m-%d %H:%M:%S')
                })
        
        return sorted(files, key=lambda x: x['time'], reverse=True)
    
    except Exception as e:
        logger.error(f"列出导出文件失败: {e}")
        return []
