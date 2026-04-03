"""
Project manager for create/load/save/delete/export flows.
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import re
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Dict, List, Optional, Tuple


logger = logging.getLogger(__name__)

PROJECTS_DIR = Path("projects")
PROJECTS_DIR.mkdir(exist_ok=True)


class ProjectManager:
    """Project storage and export helpers."""

    @staticmethod
    def _slugify(name: str) -> str:
        text = str(name or "").lower()
        text = re.sub(r"[^a-z0-9]+", "-", text)
        text = text.strip("-")
        return text or "project"

    @staticmethod
    def _safe_filename(name: str) -> str:
        cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", str(name or "")).strip()
        cleaned = cleaned.rstrip(". ")
        if not cleaned:
            return ProjectManager._slugify(name)
        return cleaned[:100]

    @staticmethod
    def create_project(
        title: str,
        genre: str,
        character_setting: str,
        world_setting: str,
        plot_idea: str,
        chapter_count: int = 50,
        owner_user_id: str = "",
        owner_email: str = "",
    ) -> Tuple[Optional[str], str]:
        if not title or not title.strip():
            return None, "项目标题不能为空"
        if not genre or not genre.strip():
            return None, "小说类型不能为空"

        normalized_title = title.strip()
        slug = ProjectManager._slugify(normalized_title)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        if slug == "project":
            title_hash = hashlib.sha1(normalized_title.encode("utf-8")).hexdigest()[:8]
            project_id = f"{slug}_{title_hash}_{timestamp}"
        else:
            project_id = f"{slug}_{timestamp}"
        project_data = {
            "id": project_id,
            "title": normalized_title,
            "genre": genre.strip(),
            "character_setting": character_setting.strip(),
            "world_setting": world_setting.strip(),
            "plot_idea": plot_idea.strip(),
            "chapter_count": chapter_count,
            "owner_user_id": owner_user_id,
            "owner_email": owner_email,
            "chapters": [],
            "outline": [],
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
        }

        try:
            project_file = PROJECTS_DIR / f"{project_id}.json"
            with open(project_file, "w", encoding="utf-8") as f:
                json.dump(project_data, f, ensure_ascii=False, indent=2)
            logger.info("[项目管理] 创建项目: %s", project_id)
            return project_id, f"项目创建成功: {project_id}"
        except Exception as exc:
            logger.error("[项目管理] 项目创建失败: %s", exc, exc_info=True)
            return None, f"项目创建失败: {exc}"

    @staticmethod
    def save_project(project_id: str, project_data: Dict) -> Tuple[bool, str]:
        try:
            if not project_data:
                return False, "项目数据为空"

            project_data["updated_at"] = datetime.now().isoformat()
            project_file = PROJECTS_DIR / f"{project_id}.json"
            with open(project_file, "w", encoding="utf-8") as f:
                json.dump(project_data, f, ensure_ascii=False, indent=2)

            logger.info("[项目管理] 项目已保存: %s", project_id)
            return True, f"项目已保存: {project_id}"
        except Exception as exc:
            logger.error("[项目管理] 项目保存失败: %s", exc, exc_info=True)
            return False, f"项目保存失败: {exc}"

    @staticmethod
    def _migrate_old_project(project_id: str, old_data: Dict) -> Dict:
        new_project_data = {
            "id": project_id,
            "title": old_data.get("title", ""),
            "genre": old_data.get("genre", ""),
            "character_setting": old_data.get("character_setting", ""),
            "world_setting": old_data.get("world_setting", ""),
            "plot_idea": old_data.get("plot_idea", ""),
            "chapter_count": old_data.get("total_chapters", len(old_data.get("chapters", []))),
            "chapters": old_data.get("chapters", []),
            "outline": old_data.get("outline", []),
            "created_at": old_data.get("created_at", ""),
            "updated_at": old_data.get("updated_at", ""),
        }
        ProjectManager.save_project(project_id, new_project_data)
        return new_project_data

    @staticmethod
    def get_project(project_id: str) -> Optional[Dict]:
        try:
            project_file = PROJECTS_DIR / f"{project_id}.json"
            if project_file.exists():
                with open(project_file, "r", encoding="utf-8") as f:
                    return json.load(f)

            project_dir = PROJECTS_DIR / project_id
            metadata_file = project_dir / "metadata.json"
            if metadata_file.exists():
                with open(metadata_file, "r", encoding="utf-8") as f:
                    old_data = json.load(f)
                logger.info("[项目管理] 发现旧格式项目，开始迁移: %s", project_id)
                return ProjectManager._migrate_old_project(project_id, old_data)

            logger.warning("[项目管理] 项目不存在: %s", project_id)
            return None
        except Exception as exc:
            logger.error("[项目管理] 项目加载失败: %s", exc, exc_info=True)
            return None

    @staticmethod
    def list_projects() -> List[Dict]:
        try:
            projects: List[Dict] = []
            if not PROJECTS_DIR.exists():
                return projects

            for project_file in PROJECTS_DIR.glob("*.json"):
                try:
                    with open(project_file, "r", encoding="utf-8") as f:
                        project_data = json.load(f)

                    chapters = project_data.get("chapters", [])
                    outline = project_data.get("outline", [])
                    completed_count = sum(1 for ch in chapters if str(ch.get("content", "")).strip())
                    total_words = sum(
                        int(ch.get("word_count") or len(str(ch.get("content", "")).strip()))
                        for ch in chapters
                    )
                    latest_chapter = next(
                        (ch for ch in reversed(chapters) if str(ch.get("content", "")).strip()),
                        chapters[-1] if chapters else {},
                    )

                    projects.append({
                        "id": project_file.stem,
                        "title": project_data.get("title", "未命名"),
                        "genre": project_data.get("genre", ""),
                        "owner_user_id": project_data.get("owner_user_id", ""),
                        "owner_email": project_data.get("owner_email", ""),
                        "created_at": project_data.get("created_at", ""),
                        "updated_at": project_data.get("updated_at", ""),
                        "total_chapters": project_data.get("chapter_count", len(chapters)),
                        "chapter_count": len(chapters),
                        "completed_chapters": completed_count,
                        "outline_count": len(outline),
                        "total_words": total_words,
                        "latest_chapter_title": latest_chapter.get("title", ""),
                        "latest_chapter_num": latest_chapter.get("num"),
                    })
                except Exception as exc:
                    logger.warning("[项目管理] 读取项目文件失败 %s: %s", project_file.name, exc)

            projects.sort(key=lambda item: item.get("updated_at", ""), reverse=True)
            return projects
        except Exception as exc:
            logger.error("[项目管理] 列出项目失败: %s", exc, exc_info=True)
            return []

    @staticmethod
    def get_project_by_title(project_title: str) -> Optional[Dict]:
        projects = ProjectManager.list_projects()
        for project in projects:
            if project.get("title") == project_title:
                return ProjectManager.get_project(project["id"])
        return None

    @staticmethod
    def delete_project(project_id: str) -> Tuple[bool, str]:
        try:
            project_file = PROJECTS_DIR / f"{project_id}.json"
            if not project_file.exists():
                return False, f"项目不存在: {project_id}"
            project_file.unlink()
            logger.info("[项目管理] 项目已删除: %s", project_id)
            return True, f"项目已删除: {project_id}"
        except Exception as exc:
            logger.error("[项目管理] 项目删除失败: %s", exc, exc_info=True)
            return False, f"项目删除失败: {exc}"

    @staticmethod
    def export_project(project_id: str, export_format: str = "json") -> Tuple[Optional[str], str]:
        try:
            project_data = ProjectManager.get_project(project_id)
            if not project_data:
                return None, f"项目不存在: {project_id}"

            export_dir = PROJECTS_DIR / "exports"
            export_dir.mkdir(exist_ok=True)

            title = project_data.get("title", "") or project_id
            genre = project_data.get("genre", "")
            created_at = project_data.get("created_at", "")
            exported_at = datetime.now().isoformat()
            safe_title = ProjectManager._safe_filename(title)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            chapters = [chapter for chapter in project_data.get("chapters", []) if chapter.get("content")]

            if export_format == "json":
                filename = f"{safe_title}_{timestamp}.json"
                filepath = export_dir / filename
                with open(filepath, "w", encoding="utf-8") as f:
                    json.dump(project_data, f, ensure_ascii=False, indent=2)
                logger.info("[项目管理] 项目已导出: %s, 格式: json", filename)
                return str(filepath), f"项目已导出: {filename}"

            if export_format in {"txt", "md"}:
                lines = [
                    f"# {title}",
                    "",
                    f"类型: {genre}",
                    f"创建时间: {created_at}",
                    f"导出时间: {exported_at}",
                    "",
                    "## 章节内容",
                    "",
                ]
                for chapter in chapters:
                    chapter_title = f"第{chapter.get('num', '')}章 {chapter.get('title', '')}".strip()
                    lines.extend([
                        f"### {chapter_title}",
                        "",
                        chapter.get("content", ""),
                        "",
                    ])

                ext = "md" if export_format == "md" else "txt"
                filename = f"{safe_title}_{timestamp}.{ext}"
                filepath = export_dir / filename
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write("\n".join(lines).strip() + "\n")
                logger.info("[项目管理] 项目已导出: %s, 格式: %s", filename, export_format)
                return str(filepath), f"项目已导出: {filename}"

            if export_format == "html":
                filename = f"{safe_title}_{timestamp}.html"
                filepath = export_dir / filename
                sections = []
                for chapter in chapters:
                    chapter_title = f"第{chapter.get('num', '')}章 {chapter.get('title', '')}".strip()
                    content = escape(str(chapter.get("content", ""))).replace("\n", "<br />\n")
                    sections.append(
                        f"<section><h2>{escape(chapter_title)}</h2><div class=\"chapter-body\">{content}</div></section>"
                    )

                html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{escape(title)}</title>
  <style>
    body {{
      max-width: 920px;
      margin: 0 auto;
      padding: 32px 20px 60px;
      font-family: "Microsoft YaHei", "PingFang SC", sans-serif;
      line-height: 1.8;
      color: #1f2937;
      background: #f8fafc;
    }}
    h1 {{
      margin-bottom: 8px;
      color: #0f172a;
    }}
    .meta {{
      color: #475569;
      margin-bottom: 28px;
    }}
    section {{
      margin-top: 28px;
      padding: 20px 22px;
      border-radius: 18px;
      background: #ffffff;
      border: 1px solid #e2e8f0;
      box-shadow: 0 12px 30px rgba(15, 23, 42, 0.06);
    }}
    h2 {{
      margin-top: 0;
      color: #0f766e;
    }}
  </style>
</head>
<body>
  <h1>{escape(title)}</h1>
  <div class="meta">
    <div>类型：{escape(genre)}</div>
    <div>创建时间：{escape(created_at)}</div>
    <div>导出时间：{escape(exported_at)}</div>
  </div>
  {''.join(sections) or '<section><p>暂无可导出的章节内容。</p></section>'}
</body>
</html>"""
                with open(filepath, "w", encoding="utf-8") as f:
                    f.write(html_content)
                logger.info("[项目管理] 项目已导出: %s, 格式: html", filename)
                return str(filepath), f"项目已导出: {filename}"

            if export_format == "docx":
                try:
                    from docx import Document
                except ImportError:
                    return None, "导出 Word 文档需要安装 python-docx: pip install python-docx"

                doc = Document()
                doc.add_heading(title, 0)
                doc.add_paragraph(f"类型: {genre}")
                doc.add_paragraph(f"创建时间: {created_at}")
                doc.add_paragraph(f"导出时间: {exported_at}")
                for chapter in chapters:
                    chapter_title = f"第{chapter.get('num', '')}章 {chapter.get('title', '')}".strip()
                    doc.add_heading(chapter_title, 1)
                    doc.add_paragraph(chapter.get("content", ""))

                filename = f"{safe_title}_{timestamp}.docx"
                filepath = export_dir / filename
                doc.save(filepath)
                logger.info("[项目管理] 项目已导出: %s, 格式: docx", filename)
                return str(filepath), f"项目已导出: {filename}"

            logger.warning("[项目管理] 不支持的导出格式: %s", export_format)
            return None, f"不支持的导出格式: {export_format}"
        except Exception as exc:
            logger.error("[项目管理] 项目导出失败: %s", exc, exc_info=True)
            return None, f"项目导出失败: {exc}"


def get_project_manager() -> ProjectManager:
    return ProjectManager()
